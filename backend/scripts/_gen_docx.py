"""One-shot generator for the .docx fixture in the debug/ folder.

Run from project root:
    backend/venv/Scripts/python.exe backend/scripts/_gen_docx.py
"""

from pathlib import Path

from docx import Document

OUT_DIR = Path(__file__).resolve().parents[2] / "debug"
OUT = OUT_DIR / "meeting_minutes.docx"


def build() -> None:
    doc = Document()
    doc.add_heading("Cascade Forge — All-Hands Meeting Minutes", level=0)

    doc.add_paragraph("Date: 12 February 2026")
    doc.add_paragraph("Location: Vancouver HQ, Boardroom A (with remote dial-in)")
    doc.add_paragraph(
        "Attendees: Maya Chen (CEO), Daniel Park (CTO), Priya Iyer (CFO), "
        "Tomás Alvarado (VP Engineering), Hannah Brookes (VP Product)."
    )

    doc.add_heading("Agenda", level=1)
    doc.add_paragraph("1. Q1 2026 progress check.", style="List Number")
    doc.add_paragraph("2. Hiring plan update.", style="List Number")
    doc.add_paragraph("3. Legacy CLI sunset proposal.", style="List Number")
    doc.add_paragraph("4. Berlin office expansion budget.", style="List Number")
    doc.add_paragraph("5. Open discussion.", style="List Number")

    doc.add_heading("Decisions", level=1)
    doc.add_paragraph(
        "Approved hiring 6 additional engineers in Q2 2026, split as 4 backend "
        "engineers on the PulseDB ingest team and 2 frontend engineers on the "
        "Helix CI web UI team.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Decided to sunset the legacy Node.js CLI by 30 June 2026. The new "
        "Rust CLI ships as a static binary in May 2026. Customer comms will "
        "go out the week of 24 February 2026.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Allocated USD 400,000 for the Berlin office expansion, covering a "
        "second floor lease, furniture, and one new conference room. Work to "
        "be completed by end of Q3 2026.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Deferred a decision on opening a fourth office in São Paulo until "
        "the Q2 2026 board review.",
        style="List Bullet",
    )

    doc.add_heading("Action items", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Owner"
    hdr[1].text = "Item"
    hdr[2].text = "Due date"

    rows = [
        ("Tomás Alvarado", "Draft customer-facing migration guide for the new CLI.", "28 February 2026"),
        ("Sofia Reyes", "Open the 6 Q2 engineering requisitions in the ATS.", "20 February 2026"),
        ("Priya Iyer", "Wire the Berlin expansion budget to the Berlin entity.", "15 March 2026"),
        ("Hannah Brookes", "Publish internal FAQ on the CLI sunset timeline.", "25 February 2026"),
        ("Daniel Park", "Schedule the Q1 2026 architecture review for the platform team.", "10 March 2026"),
    ]
    for owner, item, due in rows:
        r = table.add_row().cells
        r[0].text = owner
        r[1].text = item
        r[2].text = due

    doc.add_heading("Next meeting", level=1)
    doc.add_paragraph("11 March 2026, 10:00 PT.")

    doc.save(str(OUT))
    print(f"wrote {OUT.name}")


if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build()
