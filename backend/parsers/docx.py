import io

from docx import Document

from .base import BaseParser


class DocxParser(BaseParser):
    extension = ".docx"

    def extract_text(self, raw: bytes) -> str:
        doc = Document(io.BytesIO(raw))
        parts: list[str] = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                parts.append(" | ".join(cells))
        return "\n\n".join(parts)
